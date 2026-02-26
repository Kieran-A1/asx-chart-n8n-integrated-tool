from __future__ import annotations

import datetime as dt
import os
import re
import shutil
import subprocess
import textwrap
import time
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
from playwright.sync_api import sync_playwright

YAHOO_HOST_SUFFIXES: tuple[str, ...] = (
    "au.finance.yahoo.com",
    "finance.yahoo.com",
)

CHART_CONTAINER_SELECTORS: tuple[str, ...] = (
    "section[data-testid='qsp-chart']",
    "div[data-testid='qsp-chart']",
    "[data-testid='qsp-chart']",
    "section[data-testid='quote-chart']",
    "div[data-testid='quote-chart']",
    "[data-testid='quote-chart']",
    ".highcharts-container",
    ".highcharts-root",
    "section[id*='chart' i]",
    "section[class*='chart' i]",
)

CHART_KEYWORDS: tuple[str, ...] = (
    "1d",
    "5d",
    "1m",
    "6m",
    "ytd",
    "1y",
    "all",
    "advanced chart",
    "asx - delayed quote",
)

MIN_VALID_CHART_IMAGE_BYTES = 8_000

OUTPUT_DIR_PLACEHOLDERS: tuple[str, ...] = (
    "/path/to/local/directory",
    "path/to/local/directory",
)

# Common non-ticker words produced by natural language agent inputs.
ASX_CODE_STOPWORDS = {
    "A",
    "AN",
    "AND",
    "ASX",
    "AX",
    "ATTACH",
    "ATTACHED",
    "BODY",
    "CHART",
    "CHECK",
    "CODE",
    "COMPANY",
    "CREATE",
    "DAILY",
    "EMAIL",
    "FOR",
    "FINANCE",
    "GENERATE",
    "GRAPH",
    "IS",
    "MAIL",
    "OF",
    "PLEASE",
    "QUOTE",
    "REPORT",
    "SEND",
    "STOCK",
    "SUBJECT",
    "THE",
    "THIS",
    "TO",
    "TODAY",
    "TODAYS",
    "WITH",
    "YOUR",
    "YAHOO",
}


class AsxReportError(RuntimeError):
    """Raised when the report automation pipeline fails."""


def normalize_asx_code(asx_code: str | None) -> str | None:
    if not asx_code:
        return None

    raw = str(asx_code).strip()
    if not raw:
        return None

    upper = raw.upper()

    # Accept Yahoo style tickers like HUB.AX.
    dot_ax = re.fullmatch(r"([A-Z0-9]{1,6})\.AX", upper)
    if dot_ax and not dot_ax.group(1).isdigit():
        return dot_ax.group(1)

    # If the user/tool already passed a clean ticker code, trust it.
    if re.fullmatch(r"[A-Z0-9]{2,6}", upper) and not upper.isdigit():
        return upper

    # If a URL is provided, extract ticker from /quote/<CODE>.AX.
    url_match = re.search(r"https?://\S+", raw, flags=re.IGNORECASE)
    if url_match:
        try:
            parsed = urlparse(url_match.group(0).strip(".,;:!?\"'"))
            parts = [part for part in parsed.path.split("/") if part]

            for part in parts:
                cleaned = re.sub(r"[^A-Z0-9.]", "", part.upper())
                if cleaned.endswith(".AX"):
                    code = cleaned[:-3]
                    if 1 <= len(code) <= 6 and not code.isdigit() and code not in ASX_CODE_STOPWORDS:
                        return code

            query = parse_qs(parsed.query)
            for key in ("p", "symbol"):
                value = (query.get(key) or [""])[0].upper()
                value = re.sub(r"[^A-Z0-9.]", "", value)
                if value.endswith(".AX"):
                    code = value[:-3]
                    if 1 <= len(code) <= 6 and not code.isdigit() and code not in ASX_CODE_STOPWORDS:
                        return code
        except Exception:
            pass

    cleaned_text = re.sub(r"https?://\S+", " ", upper)
    cleaned_text = re.sub(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", " ", cleaned_text)

    for pattern in (
        r"\bASX\s*[:\-]\s*([A-Z0-9]{2,6})\b",
        r"\bASX\s+([A-Z0-9]{2,6})\b",
        r"\bFOR\s+([A-Z0-9]{2,6})\b",
        r"\b(?:CODE|TICKER|COMPANY)\s*[:\-]?\s*([A-Z0-9]{2,6})\b",
    ):
        match = re.search(pattern, cleaned_text)
        if match:
            code = re.sub(r"[^A-Z0-9]", "", match.group(1))
            if 2 <= len(code) <= 6 and not code.isdigit() and code not in ASX_CODE_STOPWORDS:
                return code

    for token in re.findall(r"\b[A-Z][A-Z0-9]{1,5}\b", cleaned_text):
        if token in ASX_CODE_STOPWORDS:
            continue
        if token.isdigit():
            continue
        return token

    return None


def build_asx_url(asx_code: str | None) -> str:
    code = (asx_code or "HUB").upper().replace(".AX", "")
    return f"https://au.finance.yahoo.com/quote/{code}.AX/"


def _safe_int_env(name: str, default: int, minimum: int = 0) -> int:
    raw = os.getenv(name, str(default)).strip()
    try:
        return max(minimum, int(raw))
    except (TypeError, ValueError):
        return max(minimum, default)


def _safe_bool_env(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "on"}


def _hostname(url: str) -> str:
    try:
        return (urlparse(url).hostname or "").lower()
    except Exception:
        return ""


def _is_asx_url(url: str) -> bool:
    host = _hostname(url)
    if not any(host == suffix or host.endswith(f".{suffix}") for suffix in YAHOO_HOST_SUFFIXES):
        return False

    try:
        path = (urlparse(url).path or "").lower()
    except Exception:
        return False

    return "/quote/" in path


def _image_looks_valid(path: Path) -> bool:
    try:
        return path.exists() and path.stat().st_size >= MIN_VALID_CHART_IMAGE_BYTES
    except Exception:
        return False


def _container_looks_like_chart(locator: Any) -> bool:
    try:
        return bool(
            locator.evaluate(
                f"""(el) => {{
                    const rect = el.getBoundingClientRect();
                    if (rect.width < 520 || rect.height < 220) return false;

                    const hasGraphic =
                        el.matches('canvas,svg,img') ||
                        !!el.querySelector('canvas,svg,img');

                    const text = (el.innerText || '').toLowerCase();
                    const hasKeyword = {list(CHART_KEYWORDS)!r}.some((token) => text.includes(token));

                    return hasGraphic || hasKeyword;
                }}"""
            )
        )
    except Exception:
        return False


def _find_chart_container(page: Any, require_keywords: bool = True) -> tuple[Any, str] | None:
    for selector in CHART_CONTAINER_SELECTORS:
        locator = page.locator(selector).first
        try:
            if locator.count() == 0:
                continue
            locator.wait_for(state="visible", timeout=500)

            if require_keywords:
                if not _container_looks_like_chart(locator):
                    continue
            else:
                box = locator.bounding_box()
                if not box:
                    continue
                if box.get("width", 0) < 520 or box.get("height", 0) < 220:
                    continue

            return locator, selector
        except Exception:
            continue

    return None


def capture_asx_graph(url: str, output_image: Path, timeout_ms: int = 45_000) -> dict[str, str]:
    output_image.parent.mkdir(parents=True, exist_ok=True)

    headless = _safe_bool_env("ASX_HEADLESS", default=False)
    watch_window_ms = _safe_int_env("ASX_WATCH_WINDOW_MS", default=20_000, minimum=1_000)
    watch_poll_ms = _safe_int_env("ASX_WATCH_POLL_MS", default=150, minimum=50)
    pre_close_wait_ms = _safe_int_env("ASX_PRE_CLOSE_WAIT_MS", default=2_000, minimum=0)

    final_url = url
    capture_selector: str | None = None

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(
            headless=headless,
            args=["--disable-blink-features=AutomationControlled"],
        )
        context = browser.new_context(
            viewport={"width": 1600, "height": 1200},
            user_agent=(
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/145.0.0.0 Safari/537.36"
            ),
            locale="en-AU",
            color_scheme="light",
            device_scale_factor=2,
            ignore_https_errors=True,
            extra_http_headers={"Accept-Language": "en-AU,en;q=0.9"},
        )

        page = context.new_page()
        try:
            page.goto(url, wait_until="domcontentloaded", timeout=timeout_ms)
            try:
                page.wait_for_load_state("networkidle", timeout=8_000)
            except PlaywrightTimeoutError:
                pass

            final_url = page.url
            if not _is_asx_url(final_url):
                raise AsxReportError(
                    f"Failed to open Yahoo Finance quote page: {url} (final URL: {final_url})"
                )

            deadline = time.monotonic() + (watch_window_ms / 1000.0)
            while time.monotonic() < deadline:
                found = _find_chart_container(page=page, require_keywords=True)
                if found is None:
                    found = _find_chart_container(page=page, require_keywords=False)

                if found is not None:
                    locator, selector = found
                    try:
                        locator.scroll_into_view_if_needed(timeout=800)
                    except Exception:
                        pass

                    locator.screenshot(path=str(output_image))
                    if _image_looks_valid(output_image):
                        capture_selector = selector
                        break

                page.wait_for_timeout(watch_poll_ms)

            if capture_selector is None:
                raise AsxReportError(
                    f"Chart did not render within {watch_window_ms}ms at {final_url}"
                )
        finally:
            try:
                if pre_close_wait_ms > 0 and not page.is_closed():
                    page.wait_for_timeout(pre_close_wait_ms)
            except Exception:
                pass
            context.close()
            browser.close()

    return {
        "source_url": url,
        "final_url": final_url,
        "image_path": str(output_image.resolve()),
        "capture_selector": capture_selector or "unavailable",
        "capture_attempt": "1",
        "capture_refreshes": "0",
    }


def create_word_report(
    title: str,
    source_url: str,
    image_path: Path,
    output_docx: Path,
) -> Path:
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Source: {source_url}")
    document.add_picture(str(image_path), width=Inches(6.5))
    document.save(str(output_docx))
    return output_docx



def _resolve_libreoffice_binary() -> str | None:
    for command in ("soffice", "libreoffice"):
        resolved = shutil.which(command)
        if resolved:
            return resolved

    mac_candidates = (
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        str(Path.home() / "Applications" / "LibreOffice.app" / "Contents" / "MacOS" / "soffice"),
    )
    for candidate in mac_candidates:
        path = Path(candidate)
        if path.exists() and os.access(path, os.X_OK):
            return str(path)

    return None


def _convert_docx_to_pdf_with_libreoffice(
    input_docx: Path,
    output_pdf: Path,
    timeout_seconds: int,
) -> str | None:
    binary = _resolve_libreoffice_binary()
    if not binary:
        return "LibreOffice binary not found"

    generated_pdf = output_pdf.parent / f"{input_docx.stem}.pdf"

    cmd = [
        binary,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_pdf.parent),
        str(input_docx),
    ]

    try:
        subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=timeout_seconds,
        )
    except Exception as exc:
        return f"LibreOffice convert failed: {exc}"

    if not generated_pdf.exists():
        return f"LibreOffice did not output: {generated_pdf}"

    try:
        if generated_pdf.stat().st_size <= 0:
            return f"LibreOffice output is empty: {generated_pdf}"
    except Exception as exc:
        return f"LibreOffice output check failed: {exc}"

    try:
        if generated_pdf.resolve() != output_pdf.resolve():
            shutil.move(str(generated_pdf), str(output_pdf))
    except Exception as exc:
        return f"LibreOffice output move failed: {exc}"

    return None


def convert_docx_to_pdf(input_docx: Path, output_pdf: Path) -> Path:
    output_pdf.parent.mkdir(parents=True, exist_ok=True)

    input_docx = input_docx.resolve()
    output_pdf = output_pdf.resolve()

    errors: list[str] = []

    def output_ready() -> bool:
        try:
            return output_pdf.exists() and output_pdf.stat().st_size > 0
        except Exception:
            return False

    pdf_engine = os.getenv("ASX_PDF_ENGINE", "auto").strip().lower()
    if not pdf_engine:
        pdf_engine = "auto"

    allow_libreoffice = pdf_engine in {"auto", "libreoffice", "soffice"}
    allow_docx2pdf = pdf_engine in {"auto", "docx2pdf", "word", "microsoft-word"}

    libreoffice_timeout_seconds = _safe_int_env(
        "ASX_LIBREOFFICE_TIMEOUT_SECONDS",
        default=120,
        minimum=30,
    )

    if allow_libreoffice:
        libreoffice_error = _convert_docx_to_pdf_with_libreoffice(
            input_docx=input_docx,
            output_pdf=output_pdf,
            timeout_seconds=libreoffice_timeout_seconds,
        )
        if libreoffice_error:
            errors.append(libreoffice_error)
        elif output_ready():
            return output_pdf

    if not allow_docx2pdf:
        detail = "; ".join(errors) if errors else f"unsupported ASX_PDF_ENGINE: {pdf_engine}"
        raise AsxReportError(f"PDF conversion failed for {input_docx}. Details: {detail}")

    # Attempt 1: direct convert to requested output path.
    try:
        convert(str(input_docx), str(output_pdf))
    except Exception as exc:
        errors.append(f"docx2pdf direct convert failed: {exc}")

    if output_ready():
        return output_pdf

    # Attempt 2: in-place conversion beside DOCX, then move.
    in_place_pdf = input_docx.with_suffix(".pdf")
    try:
        if in_place_pdf.exists() and in_place_pdf != output_pdf:
            in_place_pdf.unlink()
    except Exception:
        pass

    try:
        convert(str(input_docx))
    except Exception as exc:
        errors.append(f"docx2pdf in-place convert failed: {exc}")

    try:
        if in_place_pdf.exists() and in_place_pdf.stat().st_size > 0:
            if in_place_pdf != output_pdf:
                shutil.move(str(in_place_pdf), str(output_pdf))
    except Exception as exc:
        errors.append(f"move in-place pdf failed: {exc}")

    if output_ready():
        return output_pdf

    # Attempt 3: stage conversion in output folder then move.
    staging_dir = output_pdf.parent / ".pdf-staging"
    staging_dir.mkdir(parents=True, exist_ok=True)
    staging_docx = staging_dir / input_docx.name
    staging_pdf = staging_docx.with_suffix(".pdf")

    try:
        shutil.copy2(str(input_docx), str(staging_docx))
        if staging_pdf.exists():
            staging_pdf.unlink()
    except Exception as exc:
        errors.append(f"staging setup failed: {exc}")

    try:
        convert(str(staging_docx))
    except Exception as exc:
        errors.append(f"docx2pdf staging convert failed: {exc}")

    try:
        if staging_pdf.exists() and staging_pdf.stat().st_size > 0:
            shutil.move(str(staging_pdf), str(output_pdf))
    except Exception as exc:
        errors.append(f"staging move failed: {exc}")

    try:
        if staging_docx.exists():
            staging_docx.unlink()
    except Exception:
        pass

    if output_ready():
        return output_pdf

    detail = "; ".join(errors) if errors else "unknown error"
    raise AsxReportError(
        f"PDF conversion failed for {input_docx}. Details: {detail}"
    )


def send_email_with_mail_app(
    recipient: str,
    subject: str,
    body: str,
    attachment_path: Path,
) -> None:
    if not attachment_path.exists():
        raise AsxReportError(f"Attachment does not exist: {attachment_path}")

    attachment_abs = str(attachment_path.resolve())

    script = textwrap.dedent(
        """
        on run argv
            set subjectText to item 1 of argv
            set bodyText to item 2 of argv
            set recipientAddress to item 3 of argv
            set attachmentPath to item 4 of argv

            tell application "Mail"
                set outgoingMessage to make new outgoing message with properties {subject:subjectText, content:bodyText & return & return, visible:false}

                tell outgoingMessage
                    make new to recipient at end of to recipients with properties {address:recipientAddress}
                    tell content
                        make new attachment with properties {file name:(POSIX file attachmentPath)} at after the last paragraph
                    end tell
                    send
                end tell
            end tell
        end run
        """
    ).strip()

    subprocess.run(
        ["osascript", "-e", script, subject, body, recipient, attachment_abs],
        check=True,
    )


def _is_placeholder_output_dir(output_dir: Path) -> bool:
    raw = str(output_dir).strip().replace("\\", "/").rstrip("/").lower()
    return raw in OUTPUT_DIR_PLACEHOLDERS


def _can_write_directory(candidate: Path) -> bool:
    try:
        candidate.mkdir(parents=True, exist_ok=True)
        probe = candidate / ".asx-write-probe"
        probe.write_text("ok")
        probe.unlink(missing_ok=True)
        return True
    except Exception:
        return False


def _resolve_output_dir(output_dir: Path) -> Path:
    requested_raw = str(output_dir).strip()
    requested = Path(requested_raw or "output").expanduser()

    candidates: list[Path] = []
    if requested_raw and not _is_placeholder_output_dir(requested):
        candidates.append(requested)

    candidates.append(Path("output"))
    candidates.append(Path.home() / "Desktop" / "asx-mcp-output")
    candidates.append(Path("/tmp") / "asx-mcp-output")

    seen: set[str] = set()
    unique_candidates: list[Path] = []
    for candidate in candidates:
        key = str(candidate.resolve(strict=False))
        if key in seen:
            continue
        seen.add(key)
        unique_candidates.append(candidate)

    for candidate in unique_candidates:
        if _can_write_directory(candidate):
            return candidate

    raise AsxReportError(
        "No writable output directory available. Tried: "
        + ", ".join(str(path) for path in unique_candidates)
    )


def run_asx_report(
    asx_code: str | None,
    recipient: str,
    output_dir: Path,
    email_subject: str | None = None,
    email_body: str | None = None,
    send_email: bool = True,
) -> dict[str, str]:
    normalized_code = normalize_asx_code(asx_code)
    slug = normalized_code or "HUB"
    timestamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")

    resolved_output_dir = _resolve_output_dir(output_dir)
    image_path = resolved_output_dir / f"asx-{slug}-{timestamp}.png"
    docx_path = resolved_output_dir / f"asx-{slug}-{timestamp}.docx"
    pdf_path = resolved_output_dir / f"asx-{slug}-{timestamp}.pdf"

    source_url = build_asx_url(normalized_code)
    capture_result = capture_asx_graph(source_url, image_path)

    generated_at = dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    subject = email_subject or f"Yahoo Finance Chart Report: {slug.upper()} ({generated_at})"
    body = email_body or (
        "Attached is your Yahoo Finance chart report PDF.\n\n"
        f"Ticker: {(normalized_code or 'HUB').upper()}.AX\n"
        f"Source URL: {source_url}\n"
        f"Generated: {generated_at}"
    )

    create_word_report(
        title=subject,
        source_url=source_url,
        image_path=image_path,
        output_docx=docx_path,
    )
    convert_docx_to_pdf(docx_path, pdf_path)

    if send_email:
        send_email_with_mail_app(
            recipient=recipient,
            subject=subject,
            body=body,
            attachment_path=pdf_path,
        )

    return {
        "recipient": recipient,
        "source_url": source_url,
        "final_url": capture_result.get("final_url", source_url),
        "image_path": str(image_path.resolve()),
        "docx_path": str(docx_path.resolve()),
        "pdf_path": str(pdf_path.resolve()),
        "capture_selector": capture_result["capture_selector"],
        "capture_attempt": capture_result.get("capture_attempt", "1"),
        "capture_refreshes": capture_result.get("capture_refreshes", "0"),
        "sent_email": str(send_email).lower(),
    }
